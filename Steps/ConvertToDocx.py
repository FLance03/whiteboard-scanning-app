# Assuming input is png and txt file

import cv2 as cv
import pytesseract
import re
from docx import Document
from docx.shared import Inches

def ConvertToDocx(ListCC):
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    custom_config = '--oem 3 --psm 6'

    num = 0
    textlines = ''
    listImg = []
    for CC in ListCC:
        if CC['type'] == 'phrase':
            textlines += ' '+pytesseract.image_to_string(CC['img'], config=custom_config)
        else:
            # Save to '/CC/'
            cv.imwrite('./CC/'+str(num)+'.png', CC['img'])
            num+=1
    
    # equation to remove non-alphanumeric characters from the textlines
    textlines = re.sub(u'[^\u0020-\uD7FF\u0009\u000A\u000D\uE000-\uFFFD\U00010000-\U0010FFFF]+', '', textlines)

    # document
    document = Document()
    table = document.add_table(rows=3, cols=3, style='Table Grid')

    rowCellsA = table.rows[0].cells
    # rowCellsB = table.rows[2].cells
    rowCellsA[2].text = textlines
    cell = rowCellsA[2]
    paragraph = cell.add_paragraph()
    cell_r = paragraph.add_run()
    for x in range(num):
        cell_r.add_picture('./CC/'+str(x)+'.png')
    cellA = table.cell(2,0)
    cellB = table.cell(2,2)
    cellA.merge(cellB)

    document.save('output.docx')  
    # return document

#driver program
# input = [
#     {'img': cv.imread('images/testing.png'), 'type': 'text'},
#     {'img': cv.imread('images/testing.png'), 'type': 'nontext'},
#     {'img': cv.imread('images/testing.png'), 'type': 'text'}
# ]

# print(input[0]['type'])
# ConvertToDocx(input)